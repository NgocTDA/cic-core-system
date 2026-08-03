#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
split_master.py — Tach template tong the thanh 2 manh de rap.

  00_master_head.docx : bia -> het Heading 1 "DAC TA CHI TIET YEU CAU"
  90_master_tail.docx : tu Heading 1 "DAC TA YEU CAU PHI CHUC NANG" -> het

Phan noi dung mau o giua (Nhom chuc nang ...... / Chuc nang ......) bi bo di,
vi cho do se duoc merge.py chen cac file chuc nang vao.

Chay: python tools/split_master.py
"""
import copy
import sys
from pathlib import Path

from docx import Document

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

HEAD_END = "ĐẶC TẢ CHI TIẾT YÊU CẦU"
TAIL_START = "ĐẶC TẢ YÊU CẦU PHI CHỨC NĂNG"


def para_text(el) -> str:
    return "".join(t.text or "" for t in el.iter(f"{W}t")).strip()


def find_heading(body, text: str) -> int:
    for i, ch in enumerate(body):
        if ch.tag == f"{W}p" and para_text(ch) == text:
            return i
    raise SystemExit(f"Khong tim thay heading: {text}")


def build(src: Path, dst: Path, keep_from: int, keep_to: int) -> None:
    """Giu lai body[keep_from:keep_to+1], luon giu sectPr cuoi cung."""
    doc = Document(str(src))
    body = doc.element.body
    children = list(body)
    final_sect = children[-1] if children[-1].tag == f"{W}sectPr" else None

    for i, ch in enumerate(children):
        if ch is final_sect:
            continue
        if i < keep_from or i > keep_to:
            body.remove(ch)

    dst.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dst))
    print(f"OK  {dst}  ({len(list(Document(str(dst)).paragraphs))} doan)")


def main() -> None:
    src = Path(sys.argv[1] if len(sys.argv) > 1 else "templates/_normalized.docx")
    doc = Document(str(src))
    body = doc.element.body
    i_head_end = find_heading(body, HEAD_END)
    i_tail_start = find_heading(body, TAIL_START)
    last = len(list(body)) - 1

    build(src, Path("00_master_head.docx"), 0, i_head_end)
    build(src, Path("90_master_tail.docx"), i_tail_start, last)


if __name__ == "__main__":
    main()
