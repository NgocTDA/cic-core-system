#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Sinh file chuc nang mau theo de cuong v2.0 de chay thu pipeline."""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))

from docx import Document
from docx.shared import Pt

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
SRC = Path("templates/_normalized.docx")
BODY = "T-NoiDung"


def blank():
    doc = Document(str(SRC))
    for ch in list(doc.element.body):
        if ch.tag != f"{W}sectPr":
            doc.element.body.remove(ch)
    return doc


WIDTHS = {
    ("Hạng mục", "Nội dung"): [2400, 6955],
    ("STT", "Mã tính năng"): [500, 1900, 2400, 800, 800, 800, 2155],
    ("Mã quy tắc",): [2200, 3355, 1900, 1900],
    ("Bước",): [700, 1800, 3400, 3455],
    ("Mã luồng", "Điều kiện rẽ nhánh"): [1300, 2800, 3400, 1855],
    ("Mã luồng", "Tình huống ngoại lệ"): [1300, 2800, 3400, 1855],
    ("STT", "Tên thành phần"): [500, 1600, 1100, 800, 1800, 2000, 1555],
    ("STT", "Sự kiện / Thao tác"): [500, 2000, 1800, 3000, 2055],
    ("STT", "Mã thông báo"): [500, 2000, 1200, 3300, 2355],
    ("STT", "Loại"): [500, 1400, 2400, 1200, 3855],
    ("STT", "Nội dung vấn đề"): [500, 3800, 1800, 1400, 1855],
    ("Phiên bản",): [1100, 1400, 2200, 4655],
}


def widths_for(headers):
    for key, w in WIDTHS.items():
        if tuple(headers[:len(key)]) == key and len(w) == len(headers):
            return w
    return None


def tbl(doc, headers, rows):
    from make_child_template import set_widths
    t = doc.add_table(rows=1 + max(len(rows), 1), cols=len(headers))
    t.style = doc.styles["TableStyle3"]
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for r in c.paragraphs[0].runs:
            r.bold = True
    for ri, row in enumerate(rows, start=1):
        for ci, v in enumerate(row):
            t.rows[ri].cells[ci].text = v
    w = widths_for(headers)
    if w:
        set_widths(t, w)
    doc.add_paragraph("", style=BODY)
    return t


def label(doc, text):
    p = doc.add_paragraph(text, style=BODY)
    for r in p.runs:
        r.bold = True


def body(doc, text):
    doc.add_paragraph(text, style=BODY)


def head(doc, ma, ten, mo_ta, chinh, phu, vi_tri, tien_quyet, hau_dk, brd):
    doc.add_paragraph(f"Chức năng [{ma}] {ten}", style="Heading 3")
    doc.add_paragraph("Mô tả chung", style="Heading 4")
    tbl(doc, ["Hạng mục", "Nội dung"], [
        ["Mã chức năng", ma],
        ["Tên chức năng", ten],
        ["Mô tả chức năng", mo_ta],
        ["Tác nhân chính", chinh],
        ["Tác nhân phụ", phu],
        ["Vị trí chức năng", vi_tri],
        ["Điều kiện tiên quyết", tien_quyet],
        ["Hậu điều kiện", hau_dk],
        ["Chức năng tiền đề", "Không có"],
        ["Chức năng kế tiếp", "Không có"],
        ["Chức năng dùng chung", "CMP-GRID-001, CMP-PAGING-001"],
        ["Mã yêu cầu BRD", brd],
        ["Yêu cầu đặc thù", "Thời gian đáp ứng ≤ 2 giây với 100.000 bản ghi"],
    ])


def middle(doc, pq_rows, br_rows, luong_nv):
    doc.add_paragraph("Ma trận phân quyền", style="Heading 4")
    tbl(doc, ["STT", "Mã tính năng", "Tính năng / Thao tác",
              "ROLE-ADMIN", "ROLE-CBNV", "ROLE-LD", "Phạm vi dữ liệu"], pq_rows)

    doc.add_paragraph("Luồng màn hình", style="Heading 4")
    body(doc, "Không áp dụng")

    doc.add_paragraph("Sơ đồ trạng thái", style="Heading 4")
    body(doc, "Không áp dụng")

    doc.add_paragraph("Luồng nghiệp vụ", style="Heading 4")
    body(doc, luong_nv)

    doc.add_paragraph("Quy tắc nghiệp vụ", style="Heading 4")
    tbl(doc, ["Mã quy tắc", "Nội dung quy tắc", "Áp dụng cho",
              "Mã thông báo khi vi phạm"], br_rows)


def feature(doc, ma, ten, chinh, thay_the, ngoai_le, controls, sukien, thongbao):
    doc.add_paragraph(f"Tính năng [{ma}] {ten}", style="Heading 4")

    doc.add_paragraph("Mô tả yêu cầu", style="Heading 5")
    body(doc, f"Cho phép người dùng thực hiện {ten.lower()}. Hệ thống kiểm tra ràng buộc "
              "đầu vào theo Quy tắc chung trước khi xử lý.")

    doc.add_paragraph("Luồng xử lý", style="Heading 5")
    label(doc, "Luồng chính")
    tbl(doc, ["Bước", "Tác nhân", "Hành động", "Phản hồi của hệ thống"], chinh)
    label(doc, "Luồng thay thế")
    tbl(doc, ["Mã luồng", "Điều kiện rẽ nhánh", "Xử lý", "Quay về bước"], thay_the)
    label(doc, "Luồng ngoại lệ")
    tbl(doc, ["Mã luồng", "Tình huống ngoại lệ", "Xử lý của hệ thống",
              "Mã thông báo"], ngoai_le)

    doc.add_paragraph("Thiết kế giao diện", style="Heading 5")
    body(doc, "(Ảnh mockup)")

    doc.add_paragraph("Mô tả các thành phần trên giao diện", style="Heading 5")
    tbl(doc, ["STT", "Tên thành phần", "Loại control", "Bắt buộc",
              "Giá trị mặc định / Nguồn dữ liệu", "Ràng buộc (mã BR)",
              "Mã thông báo"], controls)

    doc.add_paragraph("Xử lý sự kiện và thao tác", style="Heading 5")
    tbl(doc, ["STT", "Sự kiện / Thao tác", "Điều kiện", "Xử lý của hệ thống",
              "Kết quả / Mã thông báo"], sukien)

    doc.add_paragraph("Thông báo", style="Heading 5")
    tbl(doc, ["STT", "Mã thông báo", "Loại", "Nội dung", "Điều kiện phát sinh"], thongbao)


def tail(doc, dulieu, vande):
    doc.add_paragraph("Dữ liệu và tích hợp", style="Heading 4")
    tbl(doc, ["STT", "Loại", "Tên đối tượng", "Chiều", "Mô tả / Ghi chú"], dulieu)

    doc.add_paragraph("Vấn đề còn mở", style="Heading 4")
    tbl(doc, ["STT", "Nội dung vấn đề", "Người quyết định", "Hạn chốt",
              "Trạng thái"], vande)

    doc.add_paragraph("Lịch sử thay đổi", style="Heading 4")
    tbl(doc, ["Phiên bản", "Ngày", "Người thực hiện", "Mô tả thay đổi"], [])


def main():
    Path("functions").mkdir(exist_ok=True)

    # ---------------- FUNC-XTH-001 Dang nhap -----------------------------
    d = blank()
    head(d, "FUNC-XTH-001", "Đăng nhập",
         "Xác thực người dùng trước khi truy cập hệ thống",
         "Người dùng CIC", "Hệ thống LDAP",
         "Màn hình khởi động", "Tài khoản đã được cấp và còn hiệu lực",
         "Phiên làm việc được tạo, ghi log truy cập", "BRD-CORE-AUTH-01")
    middle(d,
           [["1", "FEAT-XTH-001-01", "Đăng nhập bằng tài khoản nội bộ",
             "X", "X", "X", "Bản ghi của mình"]],
           [["BR-FUNC-XTH-001-001", "Sai mật khẩu 5 lần liên tiếp thì khoá 30 phút",
             "FEAT-XTH-001-01", "ERR_XTH_003"],
            ["BR-FUNC-XTH-001-002", "Mật khẩu tối thiểu 8 ký tự",
             "Trường Mật khẩu", "ERR_XTH_002"]],
           "Người dùng nhập thông tin xác thực, hệ thống kiểm tra và tạo phiên làm việc.")
    feature(d, "FEAT-XTH-001-01", "Đăng nhập bằng tài khoản nội bộ",
            [["1", "Người dùng", "Nhập tên đăng nhập và mật khẩu", "Kích hoạt nút Đăng nhập"],
             ["2", "Người dùng", "Bấm Đăng nhập", "Kiểm tra thông tin xác thực"],
             ["3", "Hệ thống", "Tạo phiên làm việc", "Chuyển đến trang chủ"]],
            [["ALT-01", "Tài khoản hết hiệu lực", "Chặn đăng nhập, hiển thị thông báo", "1"]],
            [["EXC-01", "Sai mật khẩu quá 5 lần", "Khoá tài khoản 30 phút", "ERR_XTH_003"]],
            [["1", "Tên đăng nhập", "Textbox", "Có", "Trống",
              "BR-FUNC-XTH-001-002", "ERR_XTH_001"],
             ["2", "Mật khẩu", "Password", "Có", "Trống",
              "BR-FUNC-XTH-001-002", "ERR_XTH_002"],
             ["3", "Đăng nhập", "Button", "-", "Disabled", "-", "-"]],
            [["1", "Bấm Đăng nhập", "Thông tin hợp lệ", "Tạo phiên, chuyển trang chủ",
              "SUC_XTH_001"],
             ["2", "Bấm Đăng nhập", "Sai 5 lần", "Khoá tài khoản 30 phút", "ERR_XTH_003"]],
            [["1", "ERR_XTH_001", "Inline", "Tên đăng nhập không được để trống", "Bỏ trống"],
             ["2", "ERR_XTH_003", "Modal", "Tài khoản tạm khoá 30 phút", "Sai quá 5 lần"]])
    tail(d,
         [["1", "Bảng CSDL", "TBL_NGUOI_DUNG", "Đọc", "Kiểm tra thông tin xác thực"],
          ["2", "Bảng CSDL", "TBL_NHAT_KY_TRUY_CAP", "Ghi", "Ghi log mỗi lần đăng nhập"]],
         [])
    d.save("functions/FUNC-XTH-001_Dang-nhap.docx")

    # ---------------- FUNC-XTH-002 Dang xuat -----------------------------
    d = blank()
    head(d, "FUNC-XTH-002", "Đăng xuất",
         "Kết thúc phiên làm việc an toàn", "Người dùng CIC", "Không có",
         "Thanh tiêu đề > Đăng xuất", "Đang có phiên làm việc hợp lệ",
         "Phiên bị huỷ, quay về màn hình đăng nhập", "BRD-CORE-AUTH-02")
    middle(d,
           [["1", "FEAT-XTH-002-01", "Đăng xuất chủ động",
             "X", "X", "X", "Bản ghi của mình"]],
           [["BR-FUNC-XTH-002-001", "Phải xác nhận trước khi huỷ phiên",
             "FEAT-XTH-002-01", "CONF_XTH_001"]],
           "Người dùng chủ động kết thúc phiên làm việc.")
    feature(d, "FEAT-XTH-002-01", "Đăng xuất chủ động",
            [["1", "Người dùng", "Bấm Đăng xuất", "Hiển thị hộp xác nhận"],
             ["2", "Người dùng", "Xác nhận", "Huỷ phiên, về màn hình đăng nhập"]],
            [["ALT-01", "Người dùng huỷ xác nhận", "Đóng hộp thoại, giữ phiên", "1"]],
            [["EXC-01", "Phiên đã hết hạn", "Chuyển thẳng về màn hình đăng nhập", "INF_XTH_001"]],
            [["1", "Đăng xuất", "Button", "-", "Luôn hiển thị", "-", "-"]],
            [["1", "Bấm Đăng xuất", "Đang có phiên", "Huỷ phiên", "SUC_XTH_002"]],
            [["1", "CONF_XTH_001", "Modal", "Xác nhận đăng xuất khỏi hệ thống?",
              "Bấm Đăng xuất"]])
    tail(d, [["1", "Bảng CSDL", "TBL_PHIEN_LAM_VIEC", "Ghi", "Cập nhật trạng thái phiên"]], [])
    d.save("functions/FUNC-XTH-002_Dang-xuat.docx")

    # ---------------- FUNC-NSD-001 Quan ly tai khoan ---------------------
    d = blank()
    head(d, "FUNC-NSD-001", "Quản lý tài khoản người dùng CIC",
         "Quản lý vòng đời tài khoản người dùng nội bộ",
         "Quản trị hệ thống", "Lãnh đạo đơn vị",
         "Quản trị hệ thống > Người dùng", "Đã đăng nhập với vai trò quản trị",
         "Tài khoản được tạo/cập nhật đúng trạng thái", "BRD-CORE-NSD-01")
    middle(d,
           [["1", "FEAT-NSD-001-01", "Tra cứu người dùng", "X", "X", "X", "Theo đơn vị"],
            ["2", "FEAT-NSD-001-02", "Tạo mới người dùng", "X", "", "", "Theo đơn vị"],
            ["3", "FEAT-NSD-001-02", "Phê duyệt tài khoản", "", "", "X", "Theo đơn vị"]],
           [["BR-FUNC-NSD-001-001", "Tên đăng nhập duy nhất toàn hệ thống",
             "FEAT-NSD-001-02", "ERR_NGUOIDUNG_001"],
            ["BR-FUNC-NSD-001-002", "CCCD từ 9 đến 12 ký tự, chỉ chứa số",
             "Trường CCCD", "ERR_NGUOIDUNG_002"],
            ["BR-FUNC-NSD-001-003", "Tài khoản mới luôn ở trạng thái Chờ phê duyệt",
             "FEAT-NSD-001-02", "-"]],
           "Quản trị viên tạo tài khoản, lãnh đạo phê duyệt, hệ thống kích hoạt.")
    feature(d, "FEAT-NSD-001-01", "Tra cứu người dùng",
            [["1", "Quản trị viên", "Nhập điều kiện tìm kiếm", "Kích hoạt nút Tìm kiếm"],
             ["2", "Quản trị viên", "Bấm Tìm kiếm", "Hiển thị 20 bản ghi/trang"]],
            [["ALT-01", "Không nhập điều kiện", "Hiển thị toàn bộ theo phân quyền", "2"]],
            [["EXC-01", "Không có kết quả", "Hiển thị thông báo", "INF_NGUOIDUNG_001"]],
            [["1", "Tên đăng nhập", "Textbox", "Không", "Trống", "-", "-"],
             ["2", "Trạng thái", "Dropdown", "Không", "Danh sách trạng thái", "-", "-"],
             ["3", "Tìm kiếm", "Button", "-", "-", "-", "-"]],
            [["1", "Bấm Tìm kiếm", "Có kết quả", "Hiển thị danh sách", "-"],
             ["2", "Bấm Xuất Excel", "Có kết quả", "Kết xuất file", "SUC_NGUOIDUNG_002"]],
            [["1", "INF_NGUOIDUNG_001", "Inline", "Không tìm thấy bản ghi phù hợp",
              "Không có kết quả"]])
    feature(d, "FEAT-NSD-001-02", "Tạo mới người dùng",
            [["1", "Quản trị viên", "Bấm Thêm mới", "Mở màn hình khai báo"],
             ["2", "Quản trị viên", "Nhập thông tin và bấm Lưu", "Kiểm tra ràng buộc"],
             ["3", "Hệ thống", "Tạo bản ghi", "Trạng thái Chờ phê duyệt"]],
            [["ALT-01", "Bấm Huỷ", "Đóng màn hình, không lưu", "1"]],
            [["EXC-01", "Trùng tên đăng nhập", "Chặn lưu", "ERR_NGUOIDUNG_001"]],
            [["1", "Tên đăng nhập", "Textbox", "Có", "Trống",
              "BR-FUNC-NSD-001-001", "ERR_NGUOIDUNG_001"],
             ["2", "CCCD", "Textbox", "Có", "Trống",
              "BR-FUNC-NSD-001-002", "ERR_NGUOIDUNG_002"],
             ["3", "Đơn vị", "Dropdown", "Có", "Danh mục đơn vị", "-", "-"]],
            [["1", "Bấm Lưu", "Hợp lệ", "Tạo bản ghi Chờ phê duyệt", "SUC_NGUOIDUNG_001"],
             ["2", "Bấm Lưu", "Trùng tên đăng nhập", "Chặn lưu", "ERR_NGUOIDUNG_001"]],
            [["1", "ERR_NGUOIDUNG_001", "Toast", "Tên đăng nhập đã tồn tại", "Trùng dữ liệu"]])
    tail(d,
         [["1", "Bảng CSDL", "TBL_NGUOI_DUNG", "Ghi", "Tạo và cập nhật tài khoản"],
          ["2", "API", "SSO_PROVISION", "Ra", "Đồng bộ tài khoản sang SSO"]],
         [])
    d.save("functions/FUNC-NSD-001_Quan-ly-tai-khoan.docx")

    # ---------------- FUNC-NSD-002 (draft, co loi co y) ------------------
    d = blank()
    d.add_paragraph("Phân quyền người dùng", style="Heading 2")   # sai cap
    d.add_paragraph("1.1 Mô tả", style="Heading 4")               # go tay so
    body(d, "Nội dung đang soạn.")
    d.save("functions/FUNC-NSD-002_Phan-quyen.docx")

    print("Đã sinh 4 file demo trong functions/")


if __name__ == "__main__":
    main()
