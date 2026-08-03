#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
new_function.py — Tao file dac ta chuc nang moi tu mau dung loai.

Lam giup BA 6 viec de quen:
  1. Chon dung CHILD_TEMPLATE cua loai chuc nang
  2. Xoa khoi huong dan
  3. Dien Loai / Ma / Ten chuc nang vao bang Mo ta chung
  4. Sinh ma tinh nang dau tien va dien vao tieu de khoi Tinh nang
  5. Dien ma tinh nang do vao dong dau bang Ma tran phan quyen
  6. Thay ma vao placeholder so do, va tao san file .puml khung neu loai bat buoc

Chay:
    python tools/new_function.py --ma FUNC-TCH-002 --ten "Đồng bộ tài khoản sang SSO" \\
        --loai TICHHOP --nhom GRP-TCH --ten-nhom "Tích hợp nội bộ" --owner ba02
"""
import argparse
import csv
import re
import sys
import unicodedata
from pathlib import Path

from docx import Document

sys.path.insert(0, str(Path(__file__).parent))
import outline  # noqa: E402

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
MANIFEST = Path("manifest.csv")
PUML_DIR = Path("diagrams")

PUML_SKELETON = """@startuml
title {ma} — {ten}
skinparam defaultFontName Times New Roman
skinparam defaultFontSize 13

participant CIC_CORE
participant «MÃ_HỆ_THỐNG_ĐỐI_TÁC»

CIC_CORE -> «MÃ_HỆ_THỐNG_ĐỐI_TÁC» : «bước 1»
«MÃ_HỆ_THỐNG_ĐỐI_TÁC» --> CIC_CORE : «phản hồi»
@enduml
"""


def slug(text: str) -> str:
    t = unicodedata.normalize("NFD", text)
    t = "".join(c for c in t if unicodedata.category(c) != "Mn")
    t = t.replace("đ", "d").replace("Đ", "D")
    t = re.sub(r"[^A-Za-z0-9]+", "-", t).strip("-")
    return t[:48]


def drop_guidance(doc: Document) -> None:
    """Xoa moi thu truoc Heading 3 dau tien."""
    body = doc.element.body
    for ch in list(body):
        if ch.tag == f"{W}p":
            pstyle = ch.find(f"{W}pPr/{W}pStyle")
            if pstyle is not None and pstyle.get(f"{W}val") == "Heading3":
                break
        if ch.tag != f"{W}sectPr":
            body.remove(ch)


def fill_kv(doc: Document, label: str, value: str) -> bool:
    for t in doc.tables:
        if t.rows[0].cells[0].text.strip() != "Hạng mục":
            continue
        for row in t.rows[1:]:
            if row.cells[0].text.strip() == label:
                row.cells[1].text = value
                return True
    return False


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--ma", required=True, help="ví dụ FUNC-TCH-002")
    ap.add_argument("--ten", required=True)
    ap.add_argument("--loai", required=True, choices=outline.ALL_PROFILES)
    ap.add_argument("--nhom", default="")
    ap.add_argument("--ten-nhom", default="")
    ap.add_argument("--owner", default="")
    ap.add_argument("--reviewer", default="lead.ba")
    ap.add_argument("--thu-tu-nhom", type=int, default=99)
    ap.add_argument("--thu-tu", type=int, default=99)
    ap.add_argument("--no-manifest", action="store_true")
    args = ap.parse_args()

    tpl = Path(f"templates/CHILD_TEMPLATE_{args.loai}.docx")
    if not tpl.exists():
        print(f"Chưa có {tpl} — chạy: python tools/make_child_template.py")
        return 1

    ma_feat = f"FEAT-{args.ma.replace('FUNC-', '')}-01"
    out = Path("functions") / f"{args.ma}_{slug(args.ten)}.docx"
    if out.exists():
        print(f"Đã tồn tại {out} — không ghi đè.")
        return 1

    doc = Document(str(tpl))
    drop_guidance(doc)

    # tieu de chuc nang
    for p in doc.paragraphs:
        if p.style.name == "Heading 3":
            for r in list(p.runs)[1:]:
                r._r.getparent().remove(r._r)
            p.runs[0].text = f"Chức năng [{args.ma}] {args.ten}"
            break

    fill_kv(doc, "Loại chức năng", args.loai)
    fill_kv(doc, "Mã chức năng", args.ma)
    fill_kv(doc, "Tên chức năng", args.ten)

    # tieu de khoi Tinh nang
    for p in doc.paragraphs:
        if p.style.name == "Heading 4" and p.text.startswith("Tính năng"):
            for r in list(p.runs)[1:]:
                r._r.getparent().remove(r._r)
            p.runs[0].text = f"Tính năng [{ma_feat}] «Tên tính năng»"
            break

    # dong dau ma tran phan quyen
    for t in doc.tables:
        hdr = [c.text.strip() for c in t.rows[0].cells]
        if len(hdr) > 2 and hdr[0] == "STT" and hdr[1] == "Mã tính năng":
            t.rows[1].cells[0].text = "1"
            t.rows[1].cells[1].text = ma_feat
            break

    # placeholder so do
    n_dia = 0
    for p in doc.paragraphs:
        if "[[DIAGRAM:" in p.text:
            for r in p.runs:
                r.text = r.text.replace("MÃ_CHỨC_NĂNG", args.ma)
            n_dia += 1

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"OK  {out}")
    print(f"    loại {args.loai} · tính năng đầu {ma_feat}")

    if n_dia:
        PUML_DIR.mkdir(exist_ok=True)
        puml = PUML_DIR / f"{args.ma}_seq-01.puml"
        if not puml.exists():
            puml.write_text(PUML_SKELETON.format(ma=args.ma, ten=args.ten),
                            encoding="utf-8")
            print(f"    tạo khung sơ đồ {puml}")
        else:
            print(f"    sơ đồ {puml.name} đã có, giữ nguyên")

    if not args.no_manifest and MANIFEST.exists():
        with MANIFEST.open(encoding="utf-8-sig") as f:
            fields = csv.DictReader(f).fieldnames
        with MANIFEST.open("a", encoding="utf-8", newline="") as f:
            csv.DictWriter(f, fieldnames=fields).writerow({
                "thu_tu_nhom": args.thu_tu_nhom, "ma_nhom": args.nhom,
                "ten_nhom": args.ten_nhom, "thu_tu_chuc_nang": args.thu_tu,
                "ma_chuc_nang": args.ma, "ten_chuc_nang": args.ten,
                "loai": args.loai, "file": str(out).replace("\\", "/"),
                "owner": args.owner, "reviewer": args.reviewer,
                "status": "draft", "ghi_chu": "",
            })
        print(f"    đã thêm dòng vào {MANIFEST}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
