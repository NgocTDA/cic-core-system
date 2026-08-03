#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
validate_child.py — Kiem tra file dac ta chuc nang (de cuong v3.0, 4 loai + DANHMUC).

Luat re nhanh theo dong "Loại chức năng" trong bang Mo ta chung. Doc de cuong tu
tools/outline.py nen khong bao gio lech voi template.

Chay:
    python tools/validate_child.py                       # toan bo functions/
    python tools/validate_child.py functions/FUNC-*.docx
    python tools/validate_child.py --approved f.docx      # kiem chat cho ban trinh duyet

ERROR chan merge. WARN khong chan nhung phai giai trinh.
"""
import argparse
import re
import sys
from pathlib import Path

from docx import Document

sys.path.insert(0, str(Path(__file__).parent))
import outline  # noqa: E402

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

ALLOWED_PARA_STYLES = {
    "Heading 3", "Heading 4", "Heading 5",
    "T-NoiDung", "T-Gach -", "T-Gach +",
    "Normal", "Caption", "List Paragraph",
}
ALLOWED_TABLE_STYLES = {"TableStyle3", "Table Grid"}

MANUAL_NUM = re.compile(r"^\s*(?:[IVXLC]+\.|\d+)(?:\.\d+)*[.)]?\s+\S")
MANUAL_CAPTION = re.compile(r"^\s*(Hình|Bảng|Figure|Table)\s+\d+", re.IGNORECASE)
FEATURE_HEAD = re.compile(r"^\s*Tính năng\s*\[([^\]]+)\]")
FUNC_HEAD = re.compile(r"^\s*Chức năng\s*\[([^\]]+)\]")
DIAGRAM_RE = re.compile(r"\[\[DIAGRAM:\s*([^\]]+?)\s*\]\]")

# Dang chuan cua ma: doc tu outline.py de web tool va validator khong lech nhau.
BR_CODE = re.compile(outline.CODE_PATTERNS["br"])
MSG_CODE = re.compile(outline.CODE_PATTERNS["msg"])

PUML_DIR = Path("diagrams")


def para_text(el) -> str:
    return "".join(t.text or "" for t in el.iter(f"{W}t")).strip()


def table_after(doc, heading_text: str, boundaries: set):
    """Bang dau tien nam sau heading co ten cho truoc, khong vuot qua muc ke tiep."""
    body = list(doc.element.body)
    idx = next((i for i, el in enumerate(body)
                if el.tag == f"{W}p" and para_text(el) == heading_text), None)
    if idx is None:
        return None
    for el in body[idx + 1:]:
        if el.tag == f"{W}tbl":
            for t in doc.tables:
                if t._tbl is el:
                    return t
        if el.tag == f"{W}p":
            txt = para_text(el)
            if txt in boundaries or FEATURE_HEAD.match(txt):
                return None
    return None


def data_rows(tbl):
    return [[c.text.strip() for c in r.cells] for r in tbl.rows[1:]
            if any(c.text.strip() for c in r.cells)]


def detect_profile(doc, errors: list):
    """Doc dong 'Loai chuc nang' o bang dau tien sau muc Mo ta chung."""
    all_names = set()
    for p in outline.ALL_PROFILES:
        all_names |= set(outline.h4_singleton(p))
    tbl = table_after(doc, "Mô tả chung", all_names)
    if tbl is None:
        errors.append("Không đọc được bảng ở mục “Mô tả chung” — không xác định được loại.")
        return None
    for row in tbl.rows[1:]:
        if row.cells[0].text.strip() == "Loại chức năng":
            val = row.cells[1].text.strip() if len(row.cells) > 1 else ""
            prof = outline.profile_of_name(val)
            if prof is None:
                errors.append(
                    f"Dòng “Loại chức năng” ghi “{val or '(trống)'}” — phải là một trong "
                    f"{' / '.join(outline.ALL_PROFILES)}.")
            return prof
    errors.append("Bảng “Mô tả chung” thiếu dòng “Loại chức năng”.")
    return None


def check(path: Path, approved: bool = False):
    errors, warns = [], []
    doc = Document(str(path))
    paras = list(doc.paragraphs)

    if not any(p.text.strip() for p in paras):
        return [f"{path.name}: file rỗng."], []

    heads = [(p.style.name, p.text.strip()) for p in paras
             if p.style.name.startswith("Heading")]

    # 0. Loai chuc nang -> chon bo luat -------------------------------------
    profile = detect_profile(doc, errors)
    if profile is None:
        return errors, warns
    h4_expected = outline.h4_singleton(profile)
    h5_expected = outline.h5_required(profile)
    boundaries = set(h4_expected)
    info = outline.PROFILES[profile]

    # 1. Cap heading -------------------------------------------------------
    first = heads[0] if heads else None
    if first is None:
        errors.append("Không tìm thấy heading nào. File con phải bắt đầu bằng Heading 3.")
    elif first[0] != "Heading 3":
        errors.append(f"Heading đầu tiên là '{first[0]}', phải là 'Heading 3'.")

    for st, txt in heads:
        if st in ("Heading 1", "Heading 2"):
            errors.append(f"Dùng '{st}' (chỉ master được dùng): “{txt[:50]}”")
        if st in ("Heading 6", "Heading 7", "Heading 8", "Heading 9"):
            errors.append(f"Dùng '{st}' — sâu quá 5 cấp: “{txt[:50]}”")

    h3 = [t for s, t in heads if s == "Heading 3"]
    if len(h3) > 1:
        errors.append(f"Có {len(h3)} Heading 3. Mỗi file chỉ đặc tả đúng 1 chức năng.")
    if h3 and not FUNC_HEAD.match(h3[0]):
        warns.append(f"Heading 3 nên có dạng “Chức năng [MÃ] Tên”: “{h3[0][:50]}”")
    ma_cn = FUNC_HEAD.match(h3[0]).group(1) if (h3 and FUNC_HEAD.match(h3[0])) else ""

    # 2. Style -------------------------------------------------------------
    used = {p.style.name for p in paras if p.text.strip()}
    for s in sorted(used - ALLOWED_PARA_STYLES):
        if not s.startswith("Heading"):
            errors.append(f"Style không được phép: '{s}'.")
    for t in doc.tables:
        name = t.style.name if t.style is not None else "None"
        if name not in ALLOWED_TABLE_STYLES:
            warns.append(f"Bảng dùng style '{name}', nên dùng 'TableStyle3'.")

    # 3. Danh so go tay ----------------------------------------------------
    for st, txt in heads:
        if MANUAL_NUM.match(txt):
            errors.append(f"Heading gõ tay số thứ tự: “{txt[:50]}”")
    for p in paras:
        if MANUAL_CAPTION.match(p.text):
            if not any("SEQ" in (t.text or "") for t in p._p.iter(f"{W}instrText")):
                warns.append(f"Caption gõ tay (không phải trường SEQ): “{p.text[:50]}”")

    # 4. Du muc theo de cuong CUA LOAI NAY ---------------------------------
    h4 = [t for s, t in heads if s == "Heading 4"]
    for muc in h4_expected:
        n = h4.count(muc)
        if n == 0:
            errors.append(f"[{profile}] Thiếu mục Heading 4 “{muc}”.")
        elif n > 1:
            errors.append(f"Mục “{muc}” xuất hiện {n} lần, chỉ được 1 lần.")

    # muc cua loai khac lot vao -> dau hieu dung sai mau
    other = set()
    for p in outline.ALL_PROFILES:
        if p != profile:
            other |= set(outline.h4_singleton(p)) | set(outline.h5_required(p))
    other -= set(h4_expected) | set(h5_expected)
    for st, txt in heads:
        if txt in other:
            errors.append(f"Mục “{txt}” không thuộc loại {profile} — có thể đã dùng sai mẫu.")

    features = [t for t in h4 if t.startswith("Tính năng")]
    if not features:
        errors.append("Không có mục Heading 4 “Tính năng ...”. Bắt buộc tối thiểu 1.")

    # 5. Moi khoi Tinh nang phai du muc Heading 5 --------------------------
    blocks, current = {}, None
    for st, txt in heads:
        if st == "Heading 4":
            current = txt if txt.startswith("Tính năng") else None
            if current:
                blocks[current] = []
        elif st == "Heading 5" and current:
            blocks[current].append(txt)
    for feat, subs in blocks.items():
        for muc in h5_expected:
            if muc not in subs:
                errors.append(f"“{feat[:36]}” thiếu mục Heading 5 “{muc}”.")

    # 6. Ma tran phan quyen phu het cac tinh nang --------------------------
    codes = [m.group(1) for m in (FEATURE_HEAD.match(f) for f in features) if m]
    pq = table_after(doc, "Ma trận phân quyền", boundaries)
    if pq is None:
        if features:
            warns.append("Không đọc được bảng ở mục “Ma trận phân quyền”.")
    else:
        blob = "\n".join(" ".join(r) for r in data_rows(pq))
        for code in codes:
            if code and "MÃ_TÍNH_NĂNG" not in code and code not in blob:
                errors.append(f"Tính năng '{code}' chưa có dòng trong Ma trận phân quyền.")
        hdr = " ".join(c.text for c in pq.rows[0].cells)
        if "«MÃ_VAI_TRÒ" in hdr:
            warns.append("Cột vai trò trong Ma trận phân quyền còn là placeholder.")
        if profile == "TICHHOP" and "đơn vị" not in hdr.lower():
            warns.append("Loại TICHHOP: Ma trận phân quyền phải có cột mã đơn vị / hệ thống.")

    # 7. So do trinh tu ----------------------------------------------------
    refs = []
    for p in paras:
        for m in DIAGRAM_RE.finditer(p.text):
            refs.append(m.group(1).replace("MÃ_CHỨC_NĂNG", ma_cn or "MÃ_CHỨC_NĂNG"))
    if info["require_diagram"] and not refs:
        errors.append(f"Loại {profile} bắt buộc có sơ đồ trình tự — thiếu placeholder "
                      "[[DIAGRAM: «mã»_seq-01]].")
    for name in refs:
        if "MÃ_CHỨC_NĂNG" in name:
            errors.append(f"Placeholder sơ đồ chưa thay mã: [[DIAGRAM: {name}]]")
        elif not (PUML_DIR / f"{name}.puml").exists():
            errors.append(f"Không có file {PUML_DIR}/{name}.puml cho placeholder sơ đồ.")

    # 8. Ma quy tac / ma thong bao -----------------------------------------
    qt = table_after(doc, "Quy tắc nghiệp vụ", boundaries)
    if qt is not None:
        for row in data_rows(qt):
            if row[0] and not BR_CODE.match(row[0]):
                warns.append(f"Mã quy tắc '{row[0]}' không theo dạng BR-<MÃ_CN>-nnn.")
    seen_bad = set()
    for t in doc.tables:
        for row in data_rows(t):
            for cell in row:
                for tok in re.findall(r"\b(?:ERR|WAR|INF|SUC|CONF)_[A-Z0-9_]+\b", cell):
                    if not MSG_CODE.match(tok) and tok not in seen_bad:
                        seen_bad.add(tok)
                        warns.append(f"Mã thông báo '{tok}' không theo quy ước "
                                     "[LOẠI]_[NHÓM]_[3 số].")

    # 9. Bo cuc ------------------------------------------------------------
    for p in paras:
        if p._p.find(f".//{W}br[@{W}type='page']") is not None:
            warns.append(f"Có ngắt trang thủ công: “{p.text[:40]}”")
    if len(list(doc.element.body.iter(f"{W}sectPr"))) > 1:
        warns.append("File có nhiều hơn 1 sectPr — merge.py sẽ gỡ, nhưng nên tránh.")
    direct = sum(1 for p in paras for r in p.runs if r.font.name or r.font.size)
    if direct:
        warns.append(f"Có {direct} run bị định dạng trực tiếp (đặt font/cỡ chữ tay).")

    # 10. Con sot template -------------------------------------------------
    if any("XOÁ TOÀN BỘ KHỐI NÀY" in p.text for p in paras):
        errors.append("Chưa xoá khối hướng dẫn của CHILD_TEMPLATE.")
    ph = [p.text[:40] for p in paras
          if "«" in p.text and "»" in p.text and p.style.name != "Caption"]
    if ph:
        (errors if approved else warns).append(
            f"Còn {len(ph)} placeholder «…» chưa điền, ví dụ: “{ph[0]}”")

    # 11. Kiem tra rieng cho ban trinh duyet -------------------------------
    vd = table_after(doc, "Vấn đề còn mở", boundaries)
    if vd is not None and data_rows(vd):
        (errors if approved else warns).append(
            f"Mục “Vấn đề còn mở” còn {len(data_rows(vd))} dòng chưa chốt.")

    if approved:
        mtc = table_after(doc, "Mô tả chung", boundaries)
        if mtc is not None:
            for row in mtc.rows[1:]:
                lab = row.cells[0].text.strip()
                val = row.cells[1].text.strip() if len(row.cells) > 1 else ""
                if lab and not val:
                    errors.append(f"Bảng “Mô tả chung” chưa điền: {lab}")

    return errors, warns


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("files", nargs="*")
    ap.add_argument("--approved", action="store_true")
    args = ap.parse_args()

    files = [Path(f) for f in args.files] or sorted(Path("functions").glob("*.docx"))
    if not files:
        print("Không có file nào để kiểm tra.")
        return 0

    total_err = 0
    for f in files:
        errors, warns = check(f, approved=args.approved)
        total_err += len(errors)
        status = "FAIL" if errors else ("WARN" if warns else "PASS")
        print(f"[{status}] {f.name}")
        for e in errors:
            print(f"   ERROR  {e}")
        for w in warns:
            print(f"   WARN   {w}")
    print(f"\n{len(files)} file, {total_err} lỗi chặn merge.")
    return 1 if total_err else 0


if __name__ == "__main__":
    sys.exit(main())
