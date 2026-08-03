#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
merge.py — Rap tai lieu SRS tong the tu cac file dac ta chuc nang.

    00_master_head.docx
      + (Heading 2) Nhom chuc nang A
          + functions/FUNC-...-001.docx
          + functions/FUNC-...-002.docx
      + (Heading 2) Nhom chuc nang B
          + ...
      + 90_master_tail.docx
    -> build/SRS_<he thong>_<ngay>.docx

Nguon su that ve thu tu / nhom / trang thai la manifest.csv.

Chay:
    python tools/merge.py                      # chi ghep ban da duyet
    python tools/merge.py --all                # ghep ca ban draft (de review noi bo)
    python tools/merge.py --skip-validate      # bo qua kiem tra (khong khuyen khich)
"""
import argparse
import csv
import datetime as dt
import re
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docxcompose.composer import Composer

sys.path.insert(0, str(Path(__file__).parent))
import outline  # noqa: E402
from validate_child import check as validate_check  # noqa: E402

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

HEAD = Path("00_master_head.docx")
TAIL = Path("90_master_tail.docx")
MANIFEST = Path("manifest.csv")
OUTDIR = Path("build")


def load_manifest(include_all: bool):
    with MANIFEST.open(encoding="utf-8-sig") as f:
        rows = [r for r in csv.DictReader(f) if r.get("file", "").strip()]
    if not include_all:
        rows = [r for r in rows if r["status"].strip().lower() == "approved"]
    rows.sort(key=lambda r: (int(r["thu_tu_nhom"]), int(r["thu_tu_chuc_nang"])))
    return rows


DIAGRAM_RE = re.compile(r"\[\[DIAGRAM:\s*([^\]]+?)\s*\]\]")
PNG_DIR = Path("build/diagrams")
PUML_DIR = Path("diagrams")
MAX_PIC_IN = 6.3          # be hon vung chu 6.5 inch


def diagram_title(name: str) -> str:
    """Lay dong title cua file .puml lam caption; khong co thi dung ten so do."""
    puml = PUML_DIR / f"{name}.puml"
    if puml.exists():
        for line in puml.read_text(encoding="utf-8").splitlines():
            if line.strip().lower().startswith("title "):
                return line.strip()[6:].strip()
    return name


def make_seq_caption(text: str):
    """Doan caption dung truong SEQ de Word tu danh so lai sau khi ghep."""
    p = OxmlElement("w:p")
    ppr = OxmlElement("w:pPr")
    st = OxmlElement("w:pStyle")
    st.set(qn("w:val"), "Caption")
    ppr.append(st)
    p.append(ppr)

    def run(*children):
        r = OxmlElement("w:r")
        for ch in children:
            r.append(ch)
        p.append(r)

    def t(val):
        el = OxmlElement("w:t")
        el.set(qn("xml:space"), "preserve")
        el.text = val
        return el

    def fld(kind):
        el = OxmlElement("w:fldChar")
        el.set(qn("w:fldCharType"), kind)
        return el

    run(t("Hình "))
    run(fld("begin"))
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r" SEQ Hình \* ARABIC "
    run(instr)
    run(fld("separate"))
    run(t("1"))
    run(fld("end"))
    run(t(f": {text}"))
    return p


def insert_diagrams(doc: Document, ma_chuc_nang: str) -> tuple:
    """Thay placeholder [[DIAGRAM: x]] bang anh PNG kem caption tu dong danh so."""
    from PIL import Image

    done, missing = 0, []
    for para in list(doc.paragraphs):
        m = DIAGRAM_RE.search(para.text)
        if not m:
            continue
        name = m.group(1).replace("MÃ_CHỨC_NĂNG", ma_chuc_nang)
        png = PNG_DIR / f"{name}.png"
        if not png.exists():
            missing.append(name)
            continue

        w_px, h_px = Image.open(png).size
        width_in = min(MAX_PIC_IN, w_px / 96)

        for r in list(para.runs):
            r._r.getparent().remove(r._r)
        para.style = doc.styles["T-NoiDung"]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run().add_picture(str(png), width=Inches(width_in))
        para._p.addnext(make_seq_caption(diagram_title(name)))
        done += 1
    return done, missing


def normalize_sections(doc: Document) -> Document:
    """Bo cac section break ben trong file con, giu 1 sectPr cuoi va dat = continuous.

    docxcompose can it nhat 1 sectPr de xu ly, nen khong xoa sach duoc.
    Dat continuous de noi dung chay tiep, khong sinh trang trong o cho noi.
    """
    body = doc.element.body
    for sect in list(body.iter(f"{W}sectPr")):
        parent = sect.getparent()
        if parent.tag == f"{W}pPr":            # section break giua chung
            parent.remove(sect)
            if len(parent) == 0:
                parent.getparent().remove(parent)
    doc.sections[-1].start_type = WD_SECTION.CONTINUOUS
    return doc


def git_history(file_path: str):
    """Lay lich su sua doi cua 1 file tu Git. Tra ve [] neu khong phai repo Git."""
    import subprocess
    try:
        out = subprocess.run(
            ["git", "log", "--follow", "--date=format:%d/%m/%Y",
             "--pretty=format:%ad\t%an\t%s", "--", file_path],
            capture_output=True, text=True, timeout=20, check=False)
    except (OSError, subprocess.SubprocessError):
        return []
    if out.returncode != 0 or not out.stdout.strip():
        return []
    rows = [ln.split("\t", 2) for ln in out.stdout.strip().splitlines()]
    rows.reverse()                                   # cu -> moi
    return [(f"{i}.0", d, a, s) for i, (d, a, s) in enumerate(rows, start=1)]


def fill_change_log(doc: Document, file_path: str) -> None:
    """Do Git log vao bang o muc 'Lich su thay doi' cua file con."""
    from validate_child import table_after
    bounds = set()
    for prof in outline.ALL_PROFILES:
        bounds |= set(outline.h4_singleton(prof))
    tbl = table_after(doc, "Lịch sử thay đổi", bounds)
    if tbl is None:
        return
    hist = git_history(file_path)
    if not hist:
        return
    while len(tbl.rows) > 1:
        tbl._tbl.remove(tbl.rows[-1]._tr)
    for ver, date, author, subject in hist:
        cells = tbl.add_row().cells
        for c, v in zip(cells, (ver, date, author, subject)):
            c.text = v


def mark_fields_dirty(path: Path) -> None:
    """Bat co updateFields de Word hoi cap nhat Muc luc / so hinh / so bang khi mo.

    <w:updateFields/> phai nam dung vi tri trong CT_Settings: ngay truoc
    hdrShapeDefaults / footnotePr / endnotePr / compat. Chen sai cho thi Word
    bao file loi."""
    import zipfile

    from lxml import etree

    WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    after = ("hdrShapeDefaults", "footnotePr", "endnotePr", "compat",
             "docVars", "rsids", "themeFontLang", "clrSchemeMapping")

    with zipfile.ZipFile(path) as z:
        names = z.namelist()
        data = {n: z.read(n) for n in names}

    root = etree.fromstring(data["word/settings.xml"])
    tag = f"{{{WNS}}}updateFields"
    el = root.find(tag)
    if el is None:
        el = etree.Element(tag)
        pos = len(root)
        for i, child in enumerate(root):
            if etree.QName(child).localname in after:
                pos = i
                break
        root.insert(pos, el)
    el.set(f"{{{WNS}}}val", "true")

    data["word/settings.xml"] = etree.tostring(
        root, xml_declaration=True, encoding="UTF-8", standalone=True)

    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as z:
        for n in names:
            z.writestr(n, data[n])


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--all", action="store_true", help="ghep ca ban chua duyet")
    ap.add_argument("--skip-validate", action="store_true")
    ap.add_argument("--ngat-trang", action="store_true",
                    help="moi chuc nang bat dau o trang moi")
    ap.add_argument("--render", action="store_true",
                    help="render lại sơ đồ PlantUML trước khi ghép")
    ap.add_argument("--tag", default="")
    args = ap.parse_args()

    if args.render:
        r = subprocess.run([sys.executable, "tools/render_diagrams.py"])
        if r.returncode != 0:
            print("Render sơ đồ thất bại — dừng build.")
            return 1

    rows = load_manifest(args.all)
    if not rows:
        print("manifest.csv không có dòng nào phù hợp.")
        return 1

    # --- kiem tra truoc khi ghep -----------------------------------------
    missing = [r["file"] for r in rows if not Path(r["file"]).exists()]
    if missing:
        print("Thiếu file:")
        for m in missing:
            print("   -", m)
        return 1

    if not args.skip_validate:
        failed = False
        for r in rows:
            strict = r["status"].strip().lower() == "approved"
            errors, warns = validate_check(Path(r["file"]), approved=strict)
            if errors or warns:
                print(f"[{'FAIL' if errors else 'WARN'}] {Path(r['file']).name}")
                for e in errors:
                    print("   ERROR ", e)
                for w in warns:
                    print("   WARN  ", w)
            failed |= bool(errors)
        if failed:
            print("\nCó lỗi chặn merge. Sửa xong rồi chạy lại (hoặc --skip-validate).")
            print("Nếu lỗi do dán nội dung từ nguồn khác:")
            print("    python tools/clean_child.py <file>          # xem trước")
            print("    python tools/clean_child.py <file> --apply  # dọn")
            return 1

    # --- rap --------------------------------------------------------------
    master = Document(str(HEAD))
    if args.ngat_trang:
        master.styles["Heading 3"].paragraph_format.page_break_before = True
    composer = Composer(master)

    current_group, missing_png = None, []
    for r in rows:
        if r["ma_nhom"] != current_group:
            current_group = r["ma_nhom"]
            master.add_paragraph(f"Nhóm chức năng {r['ten_nhom']}", style="Heading 2")
        child = Document(r["file"])
        fill_change_log(child, r["file"])
        n_dia, miss = insert_diagrams(child, r["ma_chuc_nang"])
        child = normalize_sections(child)
        composer.append(child)
        extra = f" · {n_dia} sơ đồ" if n_dia else ""
        print(f"  + {r['ma_chuc_nang']:<16} {r['ten_chuc_nang'][:34]:<36} "
              f"[{r.get('loai', '?'):<8}] {r['owner']}{extra}")
        for name in miss:
            print(f"      ! thiếu build/diagrams/{name}.png — placeholder giữ nguyên")
            missing_png.append(name)

    composer.append(normalize_sections(Document(str(TAIL))))

    OUTDIR.mkdir(exist_ok=True)
    stamp = dt.date.today().strftime("%Y%m%d")
    tag = f"_{args.tag}" if args.tag else ("_draft" if args.all else "")
    out = OUTDIR / f"SRS_CIC_CORE{tag}_{stamp}.docx"
    composer.save(str(out))
    mark_fields_dirty(out)

    if missing_png:
        print(f"\nCÒN {len(missing_png)} sơ đồ chưa render. Chạy:")
        print("    python tools/render_diagrams.py")
    print(f"\nOK  {out}")
    print("    Mở file trong Word, bấm Yes khi được hỏi cập nhật trường,")
    print("    hoặc Ctrl+A rồi F9 để cập nhật Mục lục và số hình/bảng.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
