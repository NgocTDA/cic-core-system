#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
make_child_template.py — Sinh templates/CHILD_TEMPLATE.docx tu tools/outline.py.

De cuong khong nam trong file nay. Muon doi de cuong thi sua tools/outline.py
roi chay lai ca 3: make_child_template.py, make_child_template_md.py, validate_child.py.

File con ke thua NGUYEN VEN styles / numbering / theme cua template tong bang
cach lay templates/_normalized.docx roi xoa sach body va dung lai bo khung.

Chay: python tools/make_child_template.py
"""
import sys
from pathlib import Path

from docx import Document
from docx.shared import Twips

sys.path.insert(0, str(Path(__file__).parent))
import outline  # noqa: E402

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
SRC = Path("templates/_normalized.docx")
DSTDIR = Path("templates")

BODY = "T-NoiDung"
BULLET = "T-Gach -"


def blank_doc() -> Document:
    doc = Document(str(SRC))
    body = doc.element.body
    for ch in list(body):
        if ch.tag != f"{W}sectPr":
            body.remove(ch)
    return doc


def set_widths(t, widths):
    """Chieu rong co dinh cho tung cot (twips).

    Phai dat ca gridCol (t.columns) va tung o (cell.width) — chi dat mot trong
    hai thi Word tu can lai theo noi dung.

    t.autofit = False da tu chen <w:tblLayout w:type="fixed"/> dung vi tri
    schema; khong tu chen tay (tblLayout phai nam truoc tblCellMar/tblLook).
    """
    assert sum(widths) == outline.USABLE
    t.autofit = False
    for col, w in zip(t.columns, widths):
        col.width = Twips(w)
    for row in t.rows:
        for cell, w in zip(row.cells, widths):
            cell.width = Twips(w)


def add_table(doc, spec):
    if spec.get("label"):
        p = doc.add_paragraph(spec["label"], style=BODY)
        for r in p.runs:
            r.bold = True
    t = doc.add_table(rows=1 + spec["rows"], cols=len(spec["headers"]))
    t.style = doc.styles["TableStyle3"]
    for i, h in enumerate(spec["headers"]):
        cell = t.rows[0].cells[i]
        cell.text = h
        for r in cell.paragraphs[0].runs:
            r.bold = True
    if spec.get("labels"):
        for i, lab in enumerate(spec["labels"]):
            t.rows[i + 1].cells[0].text = lab
    set_widths(t, spec["widths"])
    doc.add_paragraph("", style=BODY)
    return t


def note(doc, text):
    """Ghi chu huong dan — dung style Caption (11pt nghieng) chu khong dinh dang
    truc tiep, de validator khong bao 'run bi dinh dang truc tiep'."""
    doc.add_paragraph(text, style="Caption")


def render_section(doc, sec, level: int):
    doc.add_paragraph(sec["name"], style=f"Heading {level}")
    if sec.get("note"):
        note(doc, sec["note"])
    if sec.get("diagram"):
        p = doc.add_paragraph(
            outline.DIAGRAM_MARK.format(ma="MÃ_CHỨC_NĂNG"), style=BODY)
        for r in p.runs:
            r.bold = True
        doc.add_paragraph("", style=BODY)
    for tspec in sec.get("tables", []):
        add_table(doc, tspec)


def guidance(doc, profile: str):
    p = doc.add_paragraph("« HƯỚNG DẪN — XOÁ TOÀN BỘ KHỐI NÀY TRƯỚC KHI NỘP »", style=BODY)
    for r in p.runs:
        r.bold = True
    info = outline.PROFILES[profile]
    p = doc.add_paragraph(
        f"Mẫu cho loại chức năng: {profile} — {info['ten']}", style=BODY)
    for r in p.runs:
        r.bold = True
    if info.get("variant_of"):
        doc.add_paragraph(
            f"Đây là biến thể rút gọn của {info['variant_of']}, không phải một loại riêng. "
            "Hành vi CRUD chuẩn đã đặc tả một lần ở component CMP-DANHMUC-001 trong tài "
            "liệu tổng; ở đây chỉ khai báo phần riêng của danh mục này.", style=BULLET)
    if info["require_diagram"]:
        doc.add_paragraph(
            "Loại này BẮT BUỘC có sơ đồ trình tự (luồng có từ 2 hệ thống trở lên). "
            "Viết file diagrams/«mã chức năng»_seq-01.puml và giữ placeholder trong file này.",
            style=BULLET)
    for line in outline.GUIDANCE:
        doc.add_paragraph(line, style=BULLET)
    p = doc.add_paragraph("Quy ước mã:", style=BODY)
    for r in p.runs:
        r.bold = True
    for loai, dang, vd in outline.CODE_RULES:
        doc.add_paragraph(f"{loai}: {dang} — ví dụ {vd}", style=BULLET)
    doc.add_paragraph("", style=BODY)


def build(profile: str) -> Path:
    doc = blank_doc()
    guidance(doc, profile)

    doc.add_paragraph(outline.TITLE, style="Heading 3")
    before, feats, after = outline.sections(profile)

    for sec in before:
        render_section(doc, sec, 4)

    doc.add_paragraph(outline.FEATURE_TITLE, style="Heading 4")
    note(doc, outline.FEATURE_NOTE)
    for sec in feats:
        render_section(doc, sec, 5)

    for sec in after:
        render_section(doc, sec, 4)

    DSTDIR.mkdir(parents=True, exist_ok=True)
    dst = DSTDIR / f"CHILD_TEMPLATE_{profile}.docx"
    doc.save(str(dst))
    return dst


def main() -> None:
    for profile in outline.ALL_PROFILES:
        dst = build(profile)
        print(f"OK  {dst.name:<34} "
              f"{len(outline.h4_singleton(profile))} mục cấp chức năng · "
              f"{len(outline.h5_required(profile))} mục trong khối Tính năng")


if __name__ == "__main__":
    main()
